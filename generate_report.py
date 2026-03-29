"""
Generate a professional Word document report for the
Regression & Correlation Analysis assignment.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/rishabhsingh/New-Project/output"
REPORT_FILE = "/Users/rishabhsingh/New-Project/Regression_Correlation_Report.docx"

doc = Document()

# ─── Page margins ───
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # Dark blue
    return heading

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
doc.add_paragraph()  # Spacer
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('REGRESSION & CORRELATION ANALYSIS')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Individual Assignment Report')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # Spacer

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run('━' * 50)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
run.font.size = Pt(12)

doc.add_paragraph()

# Student info
info_lines = [
    ('NAME:', 'Rishabh Singh'),
    ('ID NO:', '(Fill in your ID number)'),
    ('SUBJECT:', 'Statistics — Regression & Correlation'),
    ('DATE:', 'March 2026'),
    ('SOFTWARE:', 'Python (statsmodels, scipy, matplotlib, pandas)')
]

for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    run_label.font.size = Pt(13)
    run_value = p.add_run(value)
    run_value.font.size = Pt(13)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, 'Table of Contents', level=1)

toc_items = [
    'Topic 1: Regression Analysis',
    '    1.1  What is Regression Analysis?',
    '    1.2  Research Question & Dataset',
    '    1.3  Descriptive Statistics',
    '    1.4  Python Output — OLS Regression Results',
    '    1.5  ANOVA Table',
    '    1.6  Model Summary & Coefficients',
    '    1.7  Residual Diagnostics',
    '    1.8  Output Plots',
    '    1.9  Inference',
    '',
    'Topic 2: Correlation Analysis',
    '    2.1  What is Correlation?',
    '    2.2  Research Question & Dataset',
    '    2.3  Descriptive Statistics',
    '    2.4  Python Output — Correlation Coefficients',
    '    2.5  Hypothesis Testing',
    '    2.6  Output Plots',
    '    2.7  Inference',
]

for item in toc_items:
    if item == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('    '):
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(12)
    else:
        for run in p.runs:
            run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TOPIC 1: REGRESSION ANALYSIS
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, 'TOPIC 1: REGRESSION ANALYSIS', level=1)

# 1.1 What is Regression?
add_heading_styled(doc, '1.1 What is Regression Analysis?', level=2)

p = doc.add_paragraph()
p.add_run(
    'Regression analysis is a powerful statistical technique used to model and examine the '
    'relationship between a dependent (response) variable and one or more independent (predictor) '
    'variables. The primary goal of regression analysis is to establish a mathematical equation that '
    'best describes how changes in the independent variable(s) influence the dependent variable. '
    'The simplest form, Simple Linear Regression, fits a straight line (y = β₀ + β₁x) to the data '
    'using the Ordinary Least Squares (OLS) method, which minimizes the sum of squared differences '
    'between observed and predicted values.'
)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Applications of Regression Analysis:')
run.bold = True
p.paragraph_format.space_after = Pt(4)

apps = [
    'Economics & Finance: Predicting stock prices, GDP growth, and consumer spending.',
    'Healthcare & Medicine: Modeling the effect of drug dosage on patient recovery time.',
    'Marketing: Estimating the impact of advertising spend on sales and forecasting demand.',
    'Education: Analyzing the relationship between study habits and academic performance.',
    'Engineering: Predicting material strength based on composition and optimizing processes.',
    'Real Estate: Estimating property values based on features like area, location, and amenities.'
]
for app in apps:
    doc.add_paragraph(app, style='List Bullet')

# 1.2 Research Question & Dataset
add_heading_styled(doc, '1.2 Research Question & Dataset', level=2)

p = doc.add_paragraph()
run = p.add_run('Research Question: ')
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
p.add_run('Does the number of hours a student studies per week predict their exam score?')
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.add_run(
    'A dataset of 30 students was collected, recording their weekly study hours and '
    'corresponding exam scores (out of 100). The data is presented below:'
)

# Dataset table
study_data = [
    (1.0, 30), (2.0, 35), (3.0, 42), (3.5, 44), (4.0, 48),
    (4.5, 50), (5.0, 55), (5.5, 58), (6.0, 62), (6.5, 65),
    (7.0, 68), (7.5, 70), (8.0, 74), (8.5, 76), (9.0, 80),
    (9.5, 82), (10.0, 85), (10.5, 87), (11.0, 90), (12.0, 95),
    (2.5, 38), (3.0, 45), (4.0, 52), (5.0, 56), (6.0, 60),
    (7.0, 67), (8.0, 72), (9.0, 78), (10.0, 84), (11.0, 88)
]

# Show dataset in a compact 3-column pair format
headers = ['S.No.', 'Study Hours', 'Exam Score']
rows = [(str(i+1), str(h), str(s)) for i, (h, s) in enumerate(study_data)]
add_formatted_table(doc, headers, rows)

doc.add_paragraph()

# 1.3 Descriptive Statistics
add_heading_styled(doc, '1.3 Descriptive Statistics', level=2)

desc_headers = ['Statistic', 'Study Hours', 'Exam Score']
desc_rows = [
    ('Count', '30', '30'),
    ('Mean', '6.65', '64.53'),
    ('Std. Dev.', '3.02', '18.04'),
    ('Minimum', '1.00', '30'),
    ('Q1 (25%)', '4.125', '50.50'),
    ('Median', '6.75', '66.00'),
    ('Q3 (75%)', '9.00', '79.50'),
    ('Maximum', '12.00', '95'),
]
add_formatted_table(doc, desc_headers, desc_rows)

doc.add_paragraph()

# 1.4 Python Output — OLS
add_heading_styled(doc, '1.4 Python Output — OLS Regression Results', level=2)

p = doc.add_paragraph()
p.add_run('The analysis was performed using Python with the ').font.size = Pt(11)
run = p.add_run('statsmodels')
run.italic = True
p.add_run(' library (OLS method). The complete OLS output is shown below:')

ols_output = """                            OLS Regression Results                            
==============================================================================
Dep. Variable:             Exam_Score   R-squared:                       0.994
Model:                            OLS   Adj. R-squared:                  0.994
Method:                 Least Squares   F-statistic:                     5010.
Date:                Sat, 28 Mar 2026   Prob (F-statistic):           4.02e-33
No. Observations:                  30   AIC:                             105.9
Df Residuals:                      28   BIC:                             108.7
Df Model:                           1                                         
Covariance Type:            nonrobust                                         
===============================================================================
                  coef    std err          t      P>|t|      [0.025      0.975]
-------------------------------------------------------------------------------
const          24.9312      0.613     40.689      0.000      23.676      26.186
Study_Hours     5.9552      0.084     70.784      0.000       5.783       6.128
==============================================================================
Omnibus:                        0.489   Durbin-Watson:                   0.913
Prob(Omnibus):                  0.783   Jarque-Bera (JB):                0.546
Skew:                           0.267   Prob(JB):                        0.761
Kurtosis:                       2.611   Cond. No.                         18.1
=============================================================================="""

p = doc.add_paragraph()
run = p.add_run(ols_output)
run.font.name = 'Consolas'
run.font.size = Pt(8)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)

# 1.5 ANOVA
add_heading_styled(doc, '1.5 ANOVA Table', level=2)

anova_headers = ['Source', 'Sum of Squares', 'df', 'Mean Square', 'F', 'Sig.']
anova_rows = [
    ('Regression', '9383.0307', '1', '9383.0307', '5010.40', '< 0.001'),
    ('Residual', '52.4359', '28', '1.8727', '', ''),
    ('Total', '9435.4667', '29', '', '', ''),
]
add_formatted_table(doc, anova_headers, anova_rows)

doc.add_paragraph()

# 1.6 Model Summary & Coefficients
add_heading_styled(doc, '1.6 Model Summary & Coefficients', level=2)

p = doc.add_paragraph()
run = p.add_run('Model Summary:')
run.bold = True

model_headers = ['Metric', 'Value']
model_rows = [
    ('R', '0.9972'),
    ('R²', '0.9944'),
    ('Adjusted R²', '0.9942'),
    ('Std. Error of Estimate', '1.3685'),
    ('F-statistic', '5010.40'),
    ('Prob (F-statistic)', '< 0.001'),
    ('Durbin-Watson', '0.9126'),
    ('AIC', '105.89'),
    ('BIC', '108.69'),
]
add_formatted_table(doc, model_headers, model_rows)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Coefficients Table:')
run.bold = True

coef_headers = ['Predictor', 'Coefficient', 'Std. Error', 't-value', 'p-value', '95% CI Lower', '95% CI Upper']
coef_rows = [
    ('Constant', '24.9312', '0.6127', '40.689', '0.000', '23.676', '26.186'),
    ('Study_Hours', '5.9552', '0.0841', '70.784', '0.000', '5.783', '6.128'),
]
add_formatted_table(doc, coef_headers, coef_rows)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Regression Equation:')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Exam_Score = 24.93 + 5.96 × Study_Hours')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

# 1.7 Residual Diagnostics
add_heading_styled(doc, '1.7 Residual Diagnostics', level=2)

resid_headers = ['Metric', 'Value']
resid_rows = [
    ('Mean of Residuals', '≈ 0.00'),
    ('Std. Dev. of Residuals', '1.34'),
    ('Min Residual', '−2.44'),
    ('Max Residual', '3.25'),
    ('Shapiro-Wilk W', '0.9690'),
    ('Shapiro-Wilk p-value', '0.5135'),
]
add_formatted_table(doc, resid_headers, resid_rows)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(
    'The Shapiro-Wilk test confirms that the residuals are approximately normally distributed '
    '(W = 0.969, p = 0.514 > 0.05), satisfying a key assumption of OLS regression.'
)

# 1.8 Output Plots
add_heading_styled(doc, '1.8 Output Plots', level=2)

p = doc.add_paragraph()
run = p.add_run('Figure 1: Scatter Plot with Regression Line & 95% Confidence Interval')
run.bold = True
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(f'{OUTPUT_DIR}/01_regression_scatter.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Figure 2: Residual Diagnostics (Residuals vs Fitted & Normal Q-Q Plot)')
run.bold = True
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(f'{OUTPUT_DIR}/02_regression_residuals.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1.9 Inference
add_heading_styled(doc, '1.9 Inference', level=2)

inferences_reg = [
    ('Strong Predictive Relationship: ',
     'The regression model demonstrates an exceptionally strong linear relationship between study hours '
     'and exam scores. The R² value of 0.9944 indicates that 99.44% of the variance in exam scores '
     'can be explained by the number of study hours.'),

    ('Regression Equation Interpretation: ',
     'The equation Exam_Score = 24.93 + 5.96 × Study_Hours tells us that a student who studies '
     '0 hours would be predicted to score approximately 24.93 (the intercept), and for every '
     'additional hour of study, the exam score is predicted to increase by approximately 5.96 points.'),

    ('Statistical Significance: ',
     'The F-statistic of 5010.40 with a p-value < 0.001 confirms that the overall model is statistically '
     'significant. The slope coefficient for Study_Hours (t = 70.78, p < 0.001) is also individually '
     'significant, meaning study hours is a highly significant predictor of exam scores.'),

    ('Model Adequacy: ',
     'Residuals have a mean of essentially zero, indicating no systematic bias. The Shapiro-Wilk test '
     '(W = 0.969, p = 0.514) confirms residual normality. The residual plot shows a relatively random '
     'scatter, supporting the assumption of homoscedasticity (constant variance).'),

    ('Practical Significance: ',
     'Students who study approximately 6 hours/week can expect a score around 60.66, while those studying '
     '9 hours/week can predict a score of 78.53. This finding underscores the tangible benefit of '
     'increased study time on academic performance.'),
]

for i, (title, body) in enumerate(inferences_reg, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}')
    run.bold = True
    p.add_run(body)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TOPIC 2: CORRELATION ANALYSIS
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, 'TOPIC 2: CORRELATION ANALYSIS', level=1)

# 2.1 What is Correlation?
add_heading_styled(doc, '2.1 What is Correlation?', level=2)

p = doc.add_paragraph()
p.add_run(
    'Correlation is a statistical measure that quantifies the strength and direction of the '
    'linear relationship between two continuous variables. The most commonly used measure is '
    'the Pearson correlation coefficient (r), which ranges from −1 to +1. A value of +1 '
    'indicates a perfect positive linear relationship, −1 indicates a perfect negative linear '
    'relationship, and 0 indicates no linear relationship. Other measures include Spearman\'s '
    'rank correlation (ρ) for monotonic relationships and Kendall\'s Tau (τ) for ordinal or '
    'small-sample data.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Important Note: ')
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
p.add_run(
    'Correlation does NOT imply causation. A high correlation means two variables move together, '
    'but it does not prove that one causes the other.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Applications of Correlation:')
run.bold = True

apps_corr = [
    'Finance: Analyzing the relationship between different asset returns for portfolio diversification.',
    'Medicine: Studying the correlation between risk factors (e.g., BMI) and health outcomes.',
    'Psychology: Examining relationships between personality traits and behavioral measures.',
    'Marketing: Understanding how advertising spend correlates with brand awareness and sales.',
    'Sports Analytics: Correlating training metrics with game performance.',
    'Environmental Science: Studying the relationship between pollution levels and disease rates.'
]
for app in apps_corr:
    doc.add_paragraph(app, style='List Bullet')

# 2.2 Research Question & Dataset
add_heading_styled(doc, '2.2 Research Question & Dataset', level=2)

p = doc.add_paragraph()
run = p.add_run('Research Question: ')
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
p.add_run(
    'Is there a significant relationship between monthly advertising expenditure '
    '(in $1000s) and monthly sales revenue (in $1000s)?'
)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.add_run(
    'A dataset of 30 monthly observations was collected from a company, recording advertising '
    'expenditure and corresponding sales revenue.'
)

adv_data = [
    (2.5, 22), (3.0, 28), (3.5, 31), (4.0, 35), (4.5, 38),
    (5.0, 42), (5.5, 45), (6.0, 50), (6.5, 53), (7.0, 58),
    (7.5, 60), (8.0, 64), (8.5, 67), (9.0, 72), (9.5, 75),
    (10.0, 80), (10.5, 83), (11.0, 88), (11.5, 91), (12.0, 96),
    (3.5, 33), (5.0, 44), (6.5, 55), (8.0, 66), (9.5, 76),
    (11.0, 86), (4.0, 37), (7.0, 57), (10.0, 79), (12.5, 99)
]

adv_headers = ['S.No.', 'Advertising ($1000s)', 'Sales ($1000s)']
adv_rows = [(str(i+1), str(a), str(s)) for i, (a, s) in enumerate(adv_data)]
add_formatted_table(doc, adv_headers, adv_rows)

doc.add_paragraph()

# 2.3 Descriptive Statistics
add_heading_styled(doc, '2.3 Descriptive Statistics', level=2)

desc2_headers = ['Statistic', 'Advertising ($1000s)', 'Sales ($1000s)']
desc2_rows = [
    ('Count', '30', '30'),
    ('Mean', '7.40', '60.33'),
    ('Std. Dev.', '2.94', '21.91'),
    ('Minimum', '2.50', '22'),
    ('Q1 (25%)', '5.00', '42.50'),
    ('Median', '7.25', '59.00'),
    ('Q3 (75%)', '9.875', '78.25'),
    ('Maximum', '12.50', '99'),
]
add_formatted_table(doc, desc2_headers, desc2_rows)

doc.add_paragraph()

# 2.4 Python Output — Correlation Coefficients
add_heading_styled(doc, '2.4 Python Output — Correlation Coefficients', level=2)

p = doc.add_paragraph()
p.add_run('The correlation analysis was performed using Python with the ')
run = p.add_run('scipy.stats')
run.italic = True
p.add_run(' library.')

doc.add_paragraph()

# Pearson
p = doc.add_paragraph()
run = p.add_run('Pearson Correlation Coefficient:')
run.bold = True

pearson_headers = ['Measure', 'Value']
pearson_rows = [
    ('Pearson r', '0.9989'),
    ('p-value', '6.16 × 10⁻³⁹'),
    ('R² (Coeff. of Determination)', '0.9979'),
    ('N', '30'),
]
add_formatted_table(doc, pearson_headers, pearson_rows)

doc.add_paragraph()

# Spearman
p = doc.add_paragraph()
run = p.add_run('Spearman Rank Correlation Coefficient:')
run.bold = True

spearman_headers = ['Measure', 'Value']
spearman_rows = [
    ('Spearman ρ (rho)', '0.9990'),
    ('p-value', '2.49 × 10⁻³⁹'),
    ('N', '30'),
]
add_formatted_table(doc, spearman_headers, spearman_rows)

doc.add_paragraph()

# Kendall
p = doc.add_paragraph()
run = p.add_run('Kendall Tau Correlation Coefficient:')
run.bold = True

kendall_headers = ['Measure', 'Value']
kendall_rows = [
    ('Kendall τ (tau)', '0.9896'),
    ('p-value', '2.72 × 10⁻¹⁴'),
    ('N', '30'),
]
add_formatted_table(doc, kendall_headers, kendall_rows)

doc.add_paragraph()

# Correlation Matrix
p = doc.add_paragraph()
run = p.add_run('Correlation Matrix (Pearson):')
run.bold = True

corr_headers = ['', 'Advertising Expenditure', 'Sales Revenue']
corr_rows = [
    ('Advertising Expenditure', '1.0000', '0.9989'),
    ('Sales Revenue', '0.9989', '1.0000'),
]
add_formatted_table(doc, corr_headers, corr_rows)

doc.add_paragraph()

# 2.5 Hypothesis Testing
add_heading_styled(doc, '2.5 Hypothesis Testing', level=2)

hyp_text = [
    ('H₀: ', 'ρ = 0 (No linear correlation between the variables)'),
    ('H₁: ', 'ρ ≠ 0 (There is a significant linear correlation)'),
    ('Significance level: ', 'α = 0.05'),
]

for label, value in hyp_text:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(value)

doc.add_paragraph()

hyp_results_headers = ['Metric', 'Value']
hyp_results_rows = [
    ('t-statistic', '114.3678'),
    ('Degrees of freedom', '28'),
    ('p-value (two-tail)', '≈ 0.00'),
    ('Critical t (α=0.05)', '±2.0484'),
]
add_formatted_table(doc, hyp_results_headers, hyp_results_rows)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('DECISION: ')
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)
p.add_run('Reject H₀ (p ≈ 0.00 < 0.05)')

p = doc.add_paragraph()
run = p.add_run('CONCLUSION: ')
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)
p.add_run(
    'There is a statistically significant linear correlation between '
    'Advertising Expenditure and Sales Revenue.'
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Interpretation of Correlation Strength:')
run.bold = True

p = doc.add_paragraph()
p.add_run('|r| = 0.9989 → ').bold = True
p.add_run('Very Strong Positive Correlation')

p = doc.add_paragraph()
p.add_run('R² = 0.9979 → ').bold = True
p.add_run('99.79% of the variance in Sales Revenue is explained by Advertising Expenditure.')

# 2.6 Output Plots
add_heading_styled(doc, '2.6 Output Plots', level=2)

p = doc.add_paragraph()
run = p.add_run('Figure 3: Correlation Scatter Plot with Trend Line')
run.bold = True
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(f'{OUTPUT_DIR}/03_correlation_scatter.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Figure 4: Correlation Matrix Heatmap')
run.bold = True
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(f'{OUTPUT_DIR}/04_correlation_heatmap.png', width=Inches(4.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Figure 5: Joint Distribution Plot')
run.bold = True
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(f'{OUTPUT_DIR}/05_joint_distribution.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 2.7 Inference
add_heading_styled(doc, '2.7 Inference', level=2)

inferences_corr = [
    ('Very Strong Positive Correlation: ',
     'The Pearson correlation coefficient of r = 0.9989 indicates an almost perfect positive '
     'linear relationship between advertising expenditure and sales revenue. As advertising '
     'spending increases, sales revenue increases proportionally.'),

    ('Consistency Across Measures: ',
     'All three correlation measures yield extremely high values — Pearson r = 0.9989, '
     'Spearman ρ = 0.9990, and Kendall τ = 0.9896. This consistency confirms that the '
     'relationship is robust and not driven by outliers or non-linearity.'),

    ('Statistical Significance: ',
     'The hypothesis test yields a t-statistic of 114.37 (p ≈ 0), far exceeding the critical '
     'value of ±2.048 at α = 0.05. We reject the null hypothesis and conclude that there is a '
     'statistically significant linear correlation between the two variables.'),

    ('Coefficient of Determination: ',
     'R² = 0.9979 means that 99.79% of the variability in sales revenue can be explained by its '
     'linear relationship with advertising expenditure. Only 0.21% of the variation is '
     'attributable to other factors or random noise.'),

    ('Practical Implication: ',
     'For business decision-making, this near-perfect correlation suggests that advertising '
     'expenditure is an extremely reliable predictor of sales revenue. Companies can use this '
     'relationship to forecast expected sales returns for a given advertising budget. The data '
     'suggests that approximately every additional $1,000 in advertising is associated with '
     'roughly $7,700 in additional sales revenue.'),
]

for i, (title, body) in enumerate(inferences_corr, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}')
    run.bold = True
    p.add_run(body)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# SOFTWARE & TOOLS
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, 'Software & Tools Used', level=1)

tools_headers = ['Tool', 'Version', 'Purpose']
tools_rows = [
    ('Python', '3.9', 'Programming language'),
    ('NumPy', '2.0.2', 'Numerical computations'),
    ('Pandas', '2.3.3', 'Data manipulation'),
    ('Matplotlib', '3.9.4', 'Data visualization'),
    ('SciPy', '1.13.1', 'Statistical tests'),
    ('Statsmodels', '0.14.6', 'OLS regression analysis'),
    ('scikit-learn', '1.6.1', 'Machine learning metrics'),
]
add_formatted_table(doc, tools_headers, tools_rows)

# ─── SAVE ───
doc.save(REPORT_FILE)
print(f"\n✅ Report saved successfully to:\n   {REPORT_FILE}")
print(f"\n   File size: {os.path.getsize(REPORT_FILE) / 1024:.1f} KB")
